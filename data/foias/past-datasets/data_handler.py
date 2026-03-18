#!/usr/bin/env python3
"""
Data handler for FOIA AI chat log datasets.

Parses multiple formats (HTML Copilot exports, PDF ChatGPT exports,
DOCX manual logs) into a unified Conversation/Message data model.
Includes OCR fallback for screenshot PDFs and images.

Usage:
    from data_handler import load_all_datasets, load_agency
    conversations = load_all_datasets()              # text-only (fast)
    conversations = load_all_datasets(ocr=True)      # with OCR fallback (slower)
    txdot_convos = load_agency("txdot")
"""

import json
import os
import re
import zipfile
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Optional

import fitz  # PyMuPDF
import pandas as pd
from bs4 import BeautifulSoup
from docx import Document

DATA_DIR = Path(__file__).parent

# ============================================================
# Data model
# ============================================================

@dataclass
class Message:
    role: str  # "user" or "assistant"
    content: str
    timestamp: Optional[datetime] = None


@dataclass
class Conversation:
    messages: List[Message]
    agency: str
    source_file: str
    user_name: Optional[str] = None
    user_email: Optional[str] = None
    title: Optional[str] = None
    ai_tool: Optional[str] = None  # "ChatGPT", "Copilot", "Copilot in Outlook", etc.
    date: Optional[datetime] = None
    format_type: str = ""  # "html_copilot", "pdf_chatgpt", "docx", "pdf_copilot", "policy_doc"
    metadata: dict = field(default_factory=dict)

    @property
    def num_messages(self):
        return len(self.messages)

    @property
    def num_user_messages(self):
        return sum(1 for m in self.messages if m.role == "user")

    @property
    def num_assistant_messages(self):
        return sum(1 for m in self.messages if m.role == "assistant")

    @property
    def total_chars(self):
        return sum(len(m.content) for m in self.messages)

    @property
    def user_text(self):
        return "\n".join(m.content for m in self.messages if m.role == "user")

    @property
    def assistant_text(self):
        return "\n".join(m.content for m in self.messages if m.role == "assistant")


# ============================================================
# TxDOT HTML parser (Microsoft 365 Copilot exports)
# ============================================================

def _parse_copilot_date(date_str: str) -> Optional[datetime]:
    """Parse date strings like '12/15/20255:06 PM' (no space before time)."""
    if not date_str:
        return None
    # Fix missing space between date and time
    date_str = re.sub(r'(\d{4})(\d{1,2}:\d{2})', r'\1 \2', date_str)
    for fmt in ["%m/%d/%Y %I:%M %p", "%m/%d/%Y %H:%M", "%m/%d/%Y"]:
        try:
            return datetime.strptime(date_str.strip(), fmt)
        except ValueError:
            continue
    return None


def parse_txdot_html(filepath: str) -> List[Conversation]:
    """Parse a TxDOT Microsoft 365 Copilot HTML export."""
    with open(filepath, encoding="utf-8", errors="replace") as f:
        soup = BeautifulSoup(f.read(), "html.parser")

    messages = []
    first_user = None
    first_email = None
    first_date = None
    ai_tool = "Copilot"

    for msg_div in soup.find_all("div", class_="message"):
        classes = msg_div.get("class", [])
        is_user = "message--author" in classes
        is_bot = "message--responder" in classes

        author_span = msg_div.find("span", class_="message__author")
        date_span = msg_div.find("span", class_="message__date-item-wrapper")
        content_div = msg_div.find("div", class_="message__content")

        author_text = author_span.get_text(strip=True) if author_span else ""
        date_text = date_span.get_text(strip=True) if date_span else ""
        content_text = content_div.get_text(strip=True) if content_div else ""

        if not content_text:
            continue

        # Extract email from author field (e.g. "John Doe <john@txdot.gov>")
        email_match = re.search(r'<([^>]+@[^>]+)>', author_text)
        name = re.sub(r'\s*<[^>]+>', '', author_text).strip()

        if is_user and not first_user:
            first_user = name
            first_email = email_match.group(1) if email_match else None

        if is_bot and not ai_tool:
            ai_tool = name  # e.g. "Copilot in Outlook", "Microsoft 365 Chat"

        if is_bot:
            ai_tool = name or "Copilot"

        ts = _parse_copilot_date(date_text)
        if ts and not first_date:
            first_date = ts

        messages.append(Message(
            role="user" if is_user else "assistant",
            content=content_text,
            timestamp=ts,
        ))

    if not messages:
        return []

    # Try to extract title from filename
    title = Path(filepath).stem.replace("_", " ")

    return [Conversation(
        messages=messages,
        agency="TxDOT",
        source_file=str(filepath),
        user_name=first_user,
        user_email=first_email,
        title=title,
        ai_tool=ai_tool,
        date=first_date,
        format_type="html_copilot",
    )]


# ============================================================
# PDF parser — ChatGPT text exports
# ============================================================

def _extract_pdf_text(filepath: str) -> tuple:
    """Extract text from PDF, return (full_text, has_images, is_readable)."""
    doc = fitz.open(filepath)
    full_text = ""
    total_images = 0
    for page in doc:
        full_text += page.get_text() + "\n"
        total_images += len(page.get_images())
    doc.close()

    # Check if text is meaningful (not garbled)
    printable_ratio = sum(1 for c in full_text if c.isprintable() or c.isspace()) / max(len(full_text), 1)
    is_readable = printable_ratio > 0.85 and len(full_text.strip()) > 50

    return full_text, total_images > 0, is_readable


def _ocr_pdf(filepath: str, dpi: int = 300) -> str:
    """OCR a PDF using pdf2image + pytesseract. Returns extracted text."""
    from pdf2image import convert_from_path
    import pytesseract

    pages = convert_from_path(filepath, dpi=dpi)
    all_text = []
    for page in pages:
        text = pytesseract.image_to_string(page)
        if text.strip():
            all_text.append(text)
    return "\n\n".join(all_text)


def _ocr_image(filepath: str) -> str:
    """OCR a single image file (PNG/JPG)."""
    import pytesseract
    from PIL import Image

    img = Image.open(filepath)
    return pytesseract.image_to_string(img)


def _clean_ocr_text(text: str) -> str:
    """Clean OCR artifacts from ChatGPT UI screenshots."""
    lines = text.split("\n")
    cleaned = []
    for line in lines:
        stripped = line.strip()
        # Skip browser chrome, UI elements, and common OCR noise
        skip_patterns = [
            r'^[€©®]+',                          # OCR symbol noise
            r'^[|\[\]{}]+$',                      # Pure bracket lines
            r'chat\.openai\.com',                 # URL bar
            r'^G Gmail',                          # Browser bookmarks
            r'^Microsoft Office',                 # Browser bookmarks
            r'^\+ ?(New ?chat|New Tab)',          # ChatGPT sidebar
            r'^Default \(GPT',                    # Model selector
            r'^Regenerate response',              # UI button
            r'^Free Research Preview',            # Footer
            r'^ChatGPT (May|June|July|Aug|Sep|Oct|Nov|Dec|Jan|Feb|Mar|Apr)', # Version footer
            r'^Send a message',                   # Input box placeholder
            r'^\d+/\d+$',                         # Page numbers
        ]
        if any(re.match(pat, stripped, re.IGNORECASE) for pat in skip_patterns):
            continue
        # Skip very short noise lines (1-2 chars that aren't meaningful)
        if len(stripped) <= 2 and not stripped.lower() in ('i', 'a', 'ok', 'no'):
            continue
        cleaned.append(line)
    return "\n".join(cleaned)


def _split_chatgpt_conversations(text: str) -> List[dict]:
    """Split text-based ChatGPT export into conversations with messages."""
    conversations = []

    # Split on conversation title patterns (bold headers or "user\n" at start)
    # ChatGPT exports typically have title, then alternating user/ChatGPT blocks
    lines = text.split("\n")

    current_title = None
    current_messages = []
    current_role = None
    current_content = []

    def flush_message():
        nonlocal current_role, current_content
        if current_role and current_content:
            text = "\n".join(current_content).strip()
            if text:
                current_messages.append(Message(
                    role=current_role,
                    content=text,
                ))
        current_role = None
        current_content = []

    def flush_conversation():
        nonlocal current_title, current_messages
        flush_message()
        if current_messages:
            conversations.append({
                "title": current_title,
                "messages": current_messages,
            })
        current_title = None
        current_messages = []

    for line in lines:
        stripped = line.strip()
        lower = stripped.lower()

        # Detect role markers
        if lower in ("user", "you", "you said:"):
            flush_message()
            current_role = "user"
            continue
        elif lower in ("chatgpt", "chatgpt said:", "assistant"):
            flush_message()
            current_role = "assistant"
            continue

        # Detect page numbers like "4/27" or just digits (Minneapolis IDs)
        if re.match(r'^\d+/\d+$', stripped):
            continue
        if re.match(r'^\d{6,}$', stripped):  # Minneapolis record IDs
            continue

        # If no role is set yet and we have meaningful text, it might be a title
        if current_role is None and stripped and not stripped.isdigit():
            # Check if this looks like a conversation title
            if len(stripped) < 200 and not any(stripped.lower().startswith(w) for w in
                    ["sure", "i ", "the ", "yes", "no", "here", "to ", "based"]):
                flush_conversation()
                current_title = stripped
                continue

        # Accumulate content
        if current_role:
            current_content.append(line)

    flush_conversation()
    return conversations


def parse_chatgpt_pdf(filepath: str, agency: str, ocr: bool = False) -> List[Conversation]:
    """Parse a ChatGPT PDF export. Falls back to OCR if text extraction fails and ocr=True."""
    text, has_images, is_readable = _extract_pdf_text(filepath)
    used_ocr = False

    if not is_readable:
        if not ocr:
            return [Conversation(
                messages=[],
                agency=agency,
                source_file=str(filepath),
                format_type="pdf_screenshot" if has_images else "pdf_garbled",
                metadata={"note": "Text not machine-readable; pass ocr=True to extract via OCR"},
            )]
        # OCR fallback
        text = _ocr_pdf(filepath)
        text = _clean_ocr_text(text)
        used_ocr = True
        if not text.strip():
            return [Conversation(
                messages=[], agency=agency, source_file=str(filepath),
                format_type="pdf_ocr_empty",
                metadata={"note": "OCR produced no text"},
            )]

    # Check if this is a policy document (no user/ChatGPT markers)
    lower_text = text.lower()
    has_chat_markers = any(marker in lower_text for marker in
                           ["user\n", "\nuser\n", "chatgpt\n", "\nchatgpt\n",
                            "you said", "chatgpt said"])

    if not has_chat_markers and not used_ocr:
        # Only classify as policy_doc for text-based PDFs.
        # OCR'd screenshots won't have role markers — treat them as conversations.
        return [Conversation(
            messages=[],
            agency=agency,
            source_file=str(filepath),
            format_type="policy_doc",
            title=Path(filepath).stem,
            metadata={"text_preview": text[:500], "note": "Policy/admin document, not conversation data"},
        )]

    # Extract user name from filename if possible
    stem = Path(filepath).stem
    user_name = None
    name_match = re.match(r'^([A-Z][a-z]+(?:_[A-Z][a-z]+)*)', stem)
    if name_match:
        user_name = name_match.group(1).replace("_", " ")

    fmt = "pdf_ocr" if (not is_readable and ocr) else "pdf_chatgpt"

    convos = _split_chatgpt_conversations(text)
    results = []
    for c in convos:
        results.append(Conversation(
            messages=c["messages"],
            agency=agency,
            source_file=str(filepath),
            user_name=user_name,
            title=c.get("title"),
            ai_tool="ChatGPT",
            format_type=fmt,
        ))

    # If no conversations found, return raw text as single conversation
    if not results:
        results.append(Conversation(
            messages=[Message(role="assistant", content=text.strip())],
            agency=agency,
            source_file=str(filepath),
            user_name=user_name,
            ai_tool="ChatGPT",
            format_type=fmt,
            metadata={"note": "Could not split into user/assistant messages"},
        ))

    return results


def parse_copilot_pdf(filepath: str, agency: str, ocr: bool = False) -> List[Conversation]:
    """Parse Minneapolis-style Copilot PDFs (responses only, no user prompts)."""
    text, _, is_readable = _extract_pdf_text(filepath)

    if not is_readable:
        if not ocr:
            return [Conversation(
                messages=[], agency=agency, source_file=str(filepath),
                format_type="pdf_screenshot",
                metadata={"note": "Pass ocr=True to extract via OCR"},
            )]
        text = _ocr_pdf(filepath)
        if not text.strip():
            return [Conversation(
                messages=[], agency=agency, source_file=str(filepath),
                format_type="pdf_ocr_empty",
                metadata={"note": "OCR produced no text"},
            )]

    # Split on record IDs (7+ digit numbers on their own line)
    chunks = re.split(r'\n\d{7,}\n', text)
    messages = []
    for chunk in chunks:
        chunk = chunk.strip()
        if chunk and len(chunk) > 20:
            messages.append(Message(role="assistant", content=chunk))

    return [Conversation(
        messages=messages,
        agency=agency,
        source_file=str(filepath),
        ai_tool="Copilot",
        format_type="pdf_copilot",
        metadata={"note": "Copilot responses only; user prompts not included in release"},
    )]


# ============================================================
# DOCX parser — Seattle Education manual logs
# ============================================================

def parse_docx_logs(filepath: str, agency: str) -> List[Conversation]:
    """Parse DOCX files containing manually compiled AI usage logs."""
    doc = Document(filepath)
    user_name = Path(filepath).stem.split("_")[0]  # e.g. "Armstrong" from "Armstrong_GenAI_1.docx"

    messages = []
    current_role = None
    current_content = []

    def flush():
        nonlocal current_role, current_content
        if current_role and current_content:
            text = "\n".join(current_content).strip()
            if text:
                messages.append(Message(role=current_role, content=text))
        current_role = None
        current_content = []

    for para in doc.paragraphs:
        text = para.text.strip()
        lower = text.lower()

        # Detect role markers
        if lower in ("input", "input:", "you said:", "user", "prompt", "prompt:"):
            flush()
            current_role = "user"
            continue
        elif lower in ("output", "output:", "chatgpt said:", "chatgpt", "response", "response:"):
            flush()
            current_role = "assistant"
            continue

        # Check for inline role markers
        if lower.startswith("input:") or lower.startswith("prompt:"):
            flush()
            current_role = "user"
            current_content.append(text.split(":", 1)[1].strip())
            continue
        elif lower.startswith("output:") or lower.startswith("response:") or lower.startswith("chatgpt:"):
            flush()
            current_role = "assistant"
            current_content.append(text.split(":", 1)[1].strip())
            continue

        if text and current_role:
            current_content.append(text)
        elif text and not current_role:
            # Heuristic: if we see HTML-preformatted style, try to detect roles
            if para.style and "HTML" in (para.style.name or ""):
                if "user" in lower[:20]:
                    flush()
                    current_role = "user"
                    current_content.append(text)
                elif "chatgpt" in lower[:20]:
                    flush()
                    current_role = "assistant"
                    current_content.append(text)
                elif current_role:
                    current_content.append(text)

    flush()

    if not messages:
        # Fallback: treat entire doc as unstructured text
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        messages = [Message(role="user", content=full_text)]

    return [Conversation(
        messages=messages,
        agency=agency,
        source_file=str(filepath),
        user_name=user_name,
        ai_tool="ChatGPT",
        format_type="docx",
    )]


# ============================================================
# ZIP handler — extract and parse contents
# ============================================================

def parse_zip(filepath: str, agency: str, ocr: bool = False) -> List[Conversation]:
    """Extract ZIP and parse contained files."""
    results = []
    with zipfile.ZipFile(filepath) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name_lower = info.filename.lower()
            with zf.open(info) as member:
                data = member.read()

            tmp = DATA_DIR / f"_tmp_{info.filename.replace('/', '_')}"
            tmp.write_bytes(data)
            try:
                if name_lower.endswith(".pdf"):
                    results.extend(parse_chatgpt_pdf(str(tmp), agency, ocr=ocr))
                elif name_lower.endswith((".png", ".jpg", ".jpeg")):
                    if ocr:
                        ocr_text = _ocr_image(str(tmp))
                        ocr_text = _clean_ocr_text(ocr_text)
                        if ocr_text.strip():
                            convos = _split_chatgpt_conversations(ocr_text)
                            if convos:
                                for c in convos:
                                    results.append(Conversation(
                                        messages=c["messages"], agency=agency,
                                        source_file=f"{filepath}!{info.filename}",
                                        title=c.get("title"), ai_tool="ChatGPT",
                                        format_type="image_ocr",
                                    ))
                            else:
                                results.append(Conversation(
                                    messages=[Message(role="assistant", content=ocr_text.strip())],
                                    agency=agency,
                                    source_file=f"{filepath}!{info.filename}",
                                    ai_tool="ChatGPT", format_type="image_ocr",
                                ))
                        else:
                            results.append(Conversation(
                                messages=[], agency=agency,
                                source_file=f"{filepath}!{info.filename}",
                                format_type="image_ocr_empty",
                            ))
                    else:
                        results.append(Conversation(
                            messages=[], agency=agency,
                            source_file=f"{filepath}!{info.filename}",
                            format_type="image_screenshot",
                            metadata={"note": "Screenshot image; pass ocr=True to extract"},
                        ))
            finally:
                tmp.unlink(missing_ok=True)
    return results


# ============================================================
# Agency loader configuration
# ============================================================

AGENCY_CONFIG = {
    "txdot": {
        "dir": "txdot/2026_bradford-davis_ai-usage",
        "parser": "html",
        "agency_name": "TxDOT",
        "extensions": [".html"],
    },
    "seattle-spd": {
        "dir": "seattle-spd/2023_rose-terse_chatgpt-history",
        "parser": "zip",
        "agency_name": "Seattle SPD",
        "extensions": [".zip"],
    },
    "seattle-education": {
        "dir": "seattle-education/2024_todd-feathers_genai-prompts",
        "parser": "docx",
        "agency_name": "Seattle Education",
        "extensions": [".docx"],
    },
    "spokane-pd": {
        "dir": "spokane-pd/2023-2024_rose-terse_chatgpt-history",
        "parser": "pdf_chatgpt",  # Screenshot PDFs — requires ocr=True
        "agency_name": "Spokane PD",
        "extensions": [".pdf"],
    },
    "kent-pd": {
        "dir": "kent-pd/2023_rose-terse_chatgpt-history",
        "parser": "pdf_mixed",  # Mix of text PDFs, screenshot PDFs, and a docx
        "agency_name": "Kent PD",
        "extensions": [".pdf", ".docx"],
    },
    "minneapolis-pd": {
        "dir": "minneapolis-pd/2025_joey-scott_copilot-records",
        "parser": "pdf_copilot",
        "agency_name": "Minneapolis PD",
        "extensions": [".pdf", ".PDF"],
    },
    "sec": {
        "dir": "sec/2024-2025_sungho-park_chatgpt-data",
        "parser": "pdf_policy",
        "agency_name": "SEC",
        "extensions": [".pdf"],
    },
    "cfpb": {
        "dir": "cfpb/2024_robert-delaware_chatgpt-histories",
        "parser": "pdf_policy",
        "agency_name": "CFPB",
        "extensions": [".pdf"],
    },
    "cftc": {
        "dir": "cftc/2024_robert-delaware_chatgpt-histories",
        "parser": "pdf_policy",
        "agency_name": "CFTC",
        "extensions": [".pdf"],
    },
    "fort-worth-cm": {
        "dir": "fort-worth-city-manager/2025-2026_bradford-davis_chatgpt",
        "parser": "pdf_chatgpt",
        "agency_name": "Fort Worth (City Manager)",
        "extensions": [".pdf"],
    },
    "fort-worth-ca": {
        "dir": "fort-worth-city-attorney/2025-2026_bradford-davis_chatgpt",
        "parser": "pdf_policy",  # Legal briefs, not chat data
        "agency_name": "Fort Worth (City Attorney)",
        "extensions": [".pdf"],
    },
}


def load_agency(agency_key: str, verbose: bool = False, ocr: bool = False) -> List[Conversation]:
    """Load and parse all files for a single agency.

    Args:
        ocr: If True, use OCR for screenshot PDFs and images (slower but recovers more data).
    """
    if agency_key not in AGENCY_CONFIG:
        raise ValueError(f"Unknown agency: {agency_key}. Available: {list(AGENCY_CONFIG.keys())}")

    config = AGENCY_CONFIG[agency_key]
    data_dir = DATA_DIR / config["dir"]
    agency_name = config["agency_name"]
    parser_type = config["parser"]
    extensions = config["extensions"]

    if not data_dir.exists():
        print(f"Warning: {data_dir} does not exist")
        return []

    conversations = []
    files = sorted(f for f in data_dir.iterdir()
                   if f.suffix.lower() in [e.lower() for e in extensions])

    for filepath in files:
        try:
            if parser_type == "html":
                convos = parse_txdot_html(str(filepath))
            elif parser_type == "zip":
                convos = parse_zip(str(filepath), agency_name, ocr=ocr)
            elif parser_type == "docx":
                convos = parse_docx_logs(str(filepath), agency_name)
            elif parser_type == "pdf_chatgpt":
                convos = parse_chatgpt_pdf(str(filepath), agency_name, ocr=ocr)
            elif parser_type == "pdf_copilot":
                convos = parse_copilot_pdf(str(filepath), agency_name, ocr=ocr)
            elif parser_type == "pdf_policy":
                convos = parse_chatgpt_pdf(str(filepath), agency_name, ocr=ocr)
            elif parser_type == "pdf_mixed":
                if filepath.suffix.lower() == ".docx":
                    convos = parse_docx_logs(str(filepath), agency_name)
                else:
                    convos = parse_chatgpt_pdf(str(filepath), agency_name, ocr=ocr)
            else:
                convos = []

            conversations.extend(convos)
            if verbose:
                for c in convos:
                    print(f"  {filepath.name}: {c.format_type}, {c.num_messages} messages")
        except Exception as e:
            if verbose:
                print(f"  ERROR {filepath.name}: {e}")
            conversations.append(Conversation(
                messages=[], agency=agency_name, source_file=str(filepath),
                format_type="error",
                metadata={"error": str(e)},
            ))

    return conversations


def load_all_datasets(verbose: bool = False, ocr: bool = False) -> List[Conversation]:
    """Load and parse all agency datasets.

    Args:
        ocr: If True, use OCR for screenshot PDFs and images (slower but recovers more data).
    """
    all_convos = []
    for key in AGENCY_CONFIG:
        if verbose:
            print(f"\nLoading {key}...")
        convos = load_agency(key, verbose=verbose, ocr=ocr)
        all_convos.extend(convos)
        if verbose:
            parseable = [c for c in convos if c.messages]
            print(f"  → {len(convos)} entries, {len(parseable)} with parseable messages")
    return all_convos


# ============================================================
# DataFrame conversion for analysis
# ============================================================

def conversations_to_df(conversations: List[Conversation]) -> pd.DataFrame:
    """Convert conversations to a summary DataFrame (one row per conversation)."""
    rows = []
    for c in conversations:
        rows.append({
            "agency": c.agency,
            "source_file": Path(c.source_file).name,
            "format_type": c.format_type,
            "ai_tool": c.ai_tool,
            "user_name": c.user_name,
            "user_email": c.user_email,
            "title": c.title,
            "date": c.date,
            "num_messages": c.num_messages,
            "num_user_msgs": c.num_user_messages,
            "num_asst_msgs": c.num_assistant_messages,
            "total_chars": c.total_chars,
            "has_content": c.num_messages > 0,
        })
    return pd.DataFrame(rows)


def messages_to_df(conversations: List[Conversation]) -> pd.DataFrame:
    """Convert all messages to a flat DataFrame (one row per message)."""
    rows = []
    for c in conversations:
        for i, m in enumerate(c.messages):
            rows.append({
                "agency": c.agency,
                "source_file": Path(c.source_file).name,
                "format_type": c.format_type,
                "ai_tool": c.ai_tool,
                "user_name": c.user_name,
                "user_email": c.user_email,
                "conversation_title": c.title,
                "conversation_date": c.date,
                "msg_index": i,
                "role": m.role,
                "content": m.content,
                "content_len": len(m.content),
                "timestamp": m.timestamp,
            })
    return pd.DataFrame(rows)


# ============================================================
# CLI
# ============================================================

if __name__ == "__main__":
    print("Loading all FOIA datasets...\n")
    convos = load_all_datasets(verbose=True)
    df = conversations_to_df(convos)
    print(f"\n{'='*60}")
    print(f"Total: {len(convos)} conversation entries")
    print(f"With parseable content: {df['has_content'].sum()}")
    print(f"\nBy agency:")
    print(df.groupby("agency")[["num_messages", "total_chars"]].sum().to_string())
    print(f"\nBy format:")
    print(df["format_type"].value_counts().to_string())
